from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from reportlab.pdfgen import canvas
import os

app = Flask(__name__)

# CORS Configuration to allow frontend to interact with the backend
CORS(app, resources={r"/upload": {"origins": "http://localhost:8501"}})

# Ensure the uploads folder exists
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Increase the maximum content length to 16MB for file uploads
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB limit

@app.route('/', methods=['GET'])
def home():
    return "Welcome to the Docx-to-PDF Converter API! Use the /upload route to upload a file."

@app.route('/favicon.ico')
def favicon():
    return '', 204

@app.route('/upload', methods=['POST'])
def upload_file():
    # Validate if a file is included in the request
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    # Validate file name and extension
    if not file.filename.endswith('.docx'):
        return jsonify({"error": "Invalid file format. Only .docx files are allowed."}), 400

    # Save the uploaded file to the uploads directory
    temp_doc_path = os.path.join(UPLOAD_FOLDER, file.filename)
    try:
        file.save(temp_doc_path)
    except Exception as e:
        return jsonify({"error": f"File save failed: {str(e)}"}), 500

    # Extract metadata from the docx file
    try:
        metadata = extract_metadata(temp_doc_path)
    except Exception as e:
        return jsonify({"error": f"Metadata extraction failed: {str(e)}"}), 500

    # Convert to PDF
    try:
        pdf_path = convert_to_pdf(temp_doc_path)
    except Exception as e:
        return jsonify({"error": f"PDF conversion failed: {str(e)}"}), 500

    # Return metadata and download link for the generated PDF
    download_link = f"http://localhost:5000/download/{os.path.basename(pdf_path)}"
    return jsonify({
        "message": "File uploaded and converted successfully",
        "metadata": metadata,
        "download_link": download_link
    })

def extract_metadata(file_path):
    """
    Extract metadata such as author, title, and word count from the .docx file.
    """
    try:
        doc = Document(file_path)
        metadata = {
            "author": doc.core_properties.author or "Unknown Author",
            "title": doc.core_properties.title or "Untitled",
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
        }
        return metadata
    except Exception as e:
        raise RuntimeError(f"Error reading metadata: {str(e)}")

def convert_to_pdf(file_path):
    """
    Convert the contents of a .docx file to a PDF using ReportLab.
    """
    try:
        output_pdf_path = os.path.splitext(file_path)[0] + ".pdf"
        doc = Document(file_path)
        c = canvas.Canvas(output_pdf_path)
        y = 800  # Initial Y-coordinate for the text

        for paragraph in doc.paragraphs:
            # Ensure text doesn't exceed page height
            if y < 50:  # Add a new page if content exceeds current page
                c.showPage()
                y = 800
            c.drawString(100, y, paragraph.text)
            y -= 20

        c.save()
        return output_pdf_path
    except Exception as e:
        raise RuntimeError(f"Error converting to PDF: {str(e)}")

@app.route('/download/<filename>', methods=['GET'])
def download_pdf(filename):
    """
    Serve the converted PDF file for download.
    """
    pdf_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(pdf_path):
        return jsonify({"error": "File not found"}), 404

    try:
        return send_file(pdf_path, as_attachment=True)
    except Exception as e:
        return jsonify({"error": f"Error sending file: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
