import streamlit as st
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
import os

# Streamlit UI
st.title('Word to PDF Converter')
st.write("Upload a Word document to convert it to PDF.")

# File upload
uploaded_file = st.file_uploader("Choose a .docx file", type=["docx"])

# Create a temporary folder to save uploaded files
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Handle file upload and conversion
if uploaded_file:
    # Display the uploaded file name
    st.write(f"Uploaded file: {uploaded_file.name}")

    # Save uploaded file to the temporary directory
    file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Function to extract metadata from DOCX file
    def extract_metadata(file_path):
        try:
            doc = Document(file_path)
            metadata = {
                "author": doc.core_properties.author or "Unknown Author",
                "title": doc.core_properties.title or "Untitled",
                "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
            }
            return metadata
        except Exception as e:
            raise RuntimeError(f"Error reading metadata: {str(e)}")

    # Function to convert DOCX to PDF
    def convert_to_pdf(file_path):
        try:
            output_pdf_path = os.path.splitext(file_path)[0] + ".pdf"
            doc = Document(file_path)
            pdf_buffer = BytesIO()
            c = canvas.Canvas(pdf_buffer)
            y = 800  # Initial Y-coordinate for the text

            for paragraph in doc.paragraphs:
                # Ensure text doesn't exceed page height
                if y < 50:  # Add a new page if content exceeds current page
                    c.showPage()
                    y = 800
                c.drawString(100, y, paragraph.text)
                y -= 20

            c.save()

            # Move the buffer position to the beginning
            pdf_buffer.seek(0)
            return pdf_buffer
        except Exception as e:
            raise RuntimeError(f"Error converting to PDF: {str(e)}")

    # Extract metadata from the uploaded DOCX file
    try:
        metadata = extract_metadata(file_path)
        st.write(f"**Author:** {metadata['author']}")
        st.write(f"**Title:** {metadata['title']}")
        st.write(f"**Word Count:** {metadata['word_count']}")
    except Exception as e:
        st.error(f"Error extracting metadata: {str(e)}")

    # Button to trigger the conversion
    if st.button('Upload and Convert'):
        try:
            # Convert DOCX to PDF
            pdf_file = convert_to_pdf(file_path)

            # Provide download button for the converted PDF
            st.success('File successfully uploaded and converted to PDF!')
            st.download_button(
                label="Download PDF",
                data=pdf_file,
                file_name=f"{uploaded_file.name.split('.')[0]}.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.error(f"Error during file conversion: {e}")


# import streamlit as st
# import requests

# # Streamlit UI
# st.title('Word to PDF Converter')
# st.write("Upload a Word document to convert it to PDF.")

# # File upload
# uploaded_file = st.file_uploader("Choose a .docx file", type=["docx"])

# if uploaded_file:
#     # Display the uploaded file name
#     st.write(f"Uploaded file: {uploaded_file.name}")

#     # Button to trigger upload to backend
#     if st.button('Upload and Convert'):
#         # Prepare form data to send to the Flask backend
#         file_data = {'file': uploaded_file}

#         # Send POST request to Flask backend API to handle file upload
#         backend_url = 'http://localhost:5000/upload'  # URL for the Flask API endpoint
#         response = requests.post(backend_url, files={'file': (uploaded_file.name, uploaded_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})

#         if response.status_code == 200:
#             response_json = response.json()
#             st.success('File successfully uploaded and converted!')
#             # Display metadata
#             metadata = response_json['metadata']
#             st.write(f"**Author:** {metadata['author']}")
#             st.write(f"**Title:** {metadata['title']}")
#             st.write(f"**Word Count:** {metadata['word_count']}")

#             # Provide the download link for the converted PDF
#             download_link = response_json['download_link']

#             # Streamlit's download button
#             st.download_button(
#                 label="Download PDF", 
#                 data=requests.get(download_link).content,  # Request the PDF content directly from Flask
#                 file_name=f"{uploaded_file.name.split('.')[0]}.pdf",  # Name the downloaded file
#                 mime="application/pdf"  # MIME type for PDF files
#             )

#         else:
#             st.error('Error during file conversion. Please try again later.')
